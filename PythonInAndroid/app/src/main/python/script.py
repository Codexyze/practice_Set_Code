# script.py
# import pyjokes
#
# def say_hello():
#     return pyjokes.get_joke()
# File: generate_doc.py

from docx import Document

def document(content, output_path):
    # Create document
    doc = Document()
    doc.add_heading("My Awesome Document", 0)
    doc.add_paragraph(content)

    # Save to the given path
    doc.save(output_path)
